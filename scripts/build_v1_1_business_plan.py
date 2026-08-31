#!/usr/bin/env python3
"""Build Bullion Entertainment City Business Plan V1.1 from V1.0 text + revision blocks."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path("/workspace")
SRC = Path("/home/ubuntu/.cursor/projects/workspace/agent-tools/00599e27-b8b5-4767-8679-3057bb86e510.txt")
DOCX_SRC = ROOT / "source" / "bullion_entertainment_city_business_plan_v1.0.docx"
OUT_DIR = ROOT / "金砖娱乐城商业计划书_V1.1_分章"
COMBINED_MD = ROOT / "business-plan" / "金砖娱乐城综合开发项目商业计划书_V1.1.md"
COMBINED_DOCX = ROOT / "business-plan" / "金砖娱乐城综合开发项目商业计划书_V1.1.docx"

CHAPTER_FILES = [
    ("01_执行摘要.md", "第一章"),
    ("02_项目背景与战略意义.md", "第二章"),
    ("03_项目总览与核心定位.md", "第三章"),
    ("04_建筑设计与功能规划.md", "第四章"),
    ("05_市场分析与客源定位.md", "第五章"),
    ("06_商业模式与收入模型.md", "第六章"),
    ("07_投资规划与财务预测.md", "第七章"),
    ("08_政府财政贡献.md", "第八章"),
    ("09_法律政策与合规框架.md", "第九章"),
    ("10_就业与社会责任.md", "第十章"),
    ("11_实施时间表.md", "第十一章"),
]

COVER = r'''# 金砖娱乐城综合开发项目 商业计划书（修订版 V1.1）

**机密文件 · 仅供政府审阅**

*Strictly Confidential · For Government Review Only*

| 项目名称 | 金砖娱乐城（Bullion Entertainment City） |
| :--- | :--- |
| 项目地址 | 斯里兰卡科伦坡港口城，地块编号 1-01-13 |
| 总投资规模 | 约 9.5 亿美元 |
| 主体建筑 | 金砖大厦·国家会议中心 |
| 酒店品牌 | 金砖大酒店 |
| 赌场品牌 | 澳门娱乐城 |
| 预计开业 | 2029年第二季度 |
| 文件版本 | **V1.1 · 2026年9月（修订版）** |
| 前版 | V1.0 · 2026年（政府呈报稿） |

**呈报对象：** 斯里兰卡共和国总统办公室 / 总理办公室 / 科伦坡港口城经济委员会（CPCEC）

**编制单位：** 金砖娱乐（斯里兰卡）投资有限公司

**建筑设计：** 远大集团

本文件在 V1.0 全文基础上嵌入修订条款，基准财务口径（总投资 $9.50 亿、Y3 营收约 $6.40 亿、EBITDA $2.82 亿）保持不变。V1.1 新增内容均标注【V1.1 新增】，用于补强风险披露、合同弹性、税务路径与实施前提，便于政府审阅。

未经授权不得复制、转发或对外披露。This document contains proprietary and confidential information. Unauthorized reproduction or distribution is strictly prohibited.

---

## 目录

1. 执行摘要（含1.6A宏观风险补充披露）
2. 项目背景与战略意义（含2.1.4结构性挑战）
3. 项目总览与核心定位
4. 建筑设计与功能规划（含4.4.3弹性租金条款）
5. 市场分析与客源定位（含5.3.8客群达成前提）
6. 商业模式与收入模型（含6.2.2爬坡修正情景）
7. 投资规划与财务预测（含7.2.1风险准备金）
8. 政府财政贡献（含8.2.6政策敏感性分析）
9. 法律政策与合规框架（含9.5.1 Category C税务）
10. 就业、社会责任与可持续发展
11. 实施时间表与里程碑规划（含11.6.3代理培育专项）
12. 附录：V1.0→V1.1修订摘要
'''

V11_16A = r'''
## 1.6A 宏观风险补充披露【V1.1 新增】

V1.0 第1.6节列示了基准情景回报与保守情景安全边际。为便于政府审阅，V1.1 将“不会陷入亏损困境”这一表述明确为**有前提的结论**，而非无条件保证。以下风险不改变基准财务模型，但构成决策时必须同步评估的宏观约束。

| 风险类别 | 主要表现 | 对项目的影响路径 | V1.1 对应缓释 |
| :--- | :--- | :--- | :--- |
| 主权与宏观 | IMF 计划执行波动、通胀、卢比汇率、外部融资条件收紧 | 建设成本、本地运营成本、游客信心 | 主要营收美元计价；建设期应急储备 $3,100 万；见 7.2.1 风险准备金 |
| 旅游与航空 | 国际访客不及 300 万目标；中东/印度航线运力不足 | 酒店入住率、主大厅客流、峰会配套 | 客群六源对冲；见 5.3.8 达成前提 |
| 牌照与时序 | GRA 牌照或 CPCEC 授权晚于 2027 年 9 月目标 | 赌场开业推迟，Y1 租金与 GGR 后移 | 酒店 2029 年 3 月先行开业；见第十一章缓冲 |
| 监管与 FATF | 国家互评估结果影响博彩业国际清算与银行开户 | VIP 资金通道、代理展业 | 主动按 FATF 标准建设 AML；见第九章 |
| 竞争 | City of Dreams 先发、后续牌照增加 | VIP 获客成本上升、上座率低于 70% | 错位定位；见 6.2.2 爬坡修正情景 |
| 招标厅履约 | 运营商延期、减租或退租 | $1.20 亿固定租金不能按面值入账 | 见 4.4.3 弹性租金；保证金与托管 |
| 跨境客源政策 | 印度、中国等地出境支付/签证收紧 | 泥码规模与访客结构偏离 5.3 节假设 | 海湾 + 金砖 + 医养多源；代理培育见 11.6.3 |
| 社会许可 | 宗教与公众对博彩的关切 | 审批节奏、本地就业承诺加码 | 负责任博彩 + 73.8% 本地用工（第十章） |

**对 1.6 节保守情景的限定：** 净利润约 $4,200 万的保守情景，以 12 个招标厅按约支付年租金 $1.20 亿为前提。若开业当年仅 8–10 个厅按面值计租、其余厅适用爬坡折扣或空置，则该安全垫按比例收窄。V1.1 因此增加弹性租金条款与风险准备金，使下行情景可管理，而不是把固定租金视为无风险国债。

**不纳入基准模型、但需向政府单独说明的事项：** 港口城 Primary BSI Category C（最低投资 $10 亿）与本项目现行总投资 $9.50 亿存在 $5,000 万缺口。财务模型仍按企业所得税 20% 计提，**未提前计入免税期**。税务路径详见 9.5.1。
'''

V11_214 = r'''
### 2.1.4 结构性挑战【V1.1 新增】

第 2.1.1–2.1.3 节描述了复苏窗口。窗口同时意味着约束。以下结构性挑战不否定项目战略价值，但决定实施节奏与合同必须预留弹性。

**财政改革与社会承受力并行。** IMF 要求税收/GDP 由 7.3% 提升至 15%，博彩综合体是高效税基，也可能成为公共舆论焦点。项目必须把财政贡献（第八章）与负责任博彩、本地就业、社区基金（第十章）绑定呈报，避免被理解为“只抽税、不担责”。

**能源、市政与口岸能力。** 港口城填海区电力、淡水、污水与清关峰值能否匹配 325,400 m² 超高端综合体，是建设期关键外部依赖。屋顶航空接驳还受空域许可约束。上述事项列入 CPCEC 综合授权与开工前条件，而非开业后补救。

**人才厚度。** 直接用工 4,100 人、本地占比 73.8% 的目标，与科伦坡高端酒店及赌场专业人力供给存在缺口。荷官、监察、国际礼宾、同传、再生医学医护均需提前 18–24 个月培养。V1.1 将贵宾代理与关键岗位培育单列（11.6.3），避免把“开业即满编”写成可自动实现的结果。

**双机场与 vis-à-vis 交通。** 孟买 2.5 小时、迪拜 4 小时是地理事实；VIP 闭环体验依赖汉班托塔专属飞行器与科伦坡口岸的稳定通关。若航权、签证或公务机手续未同步，5.3 节海湾客群 30–35% 的结构难以按期兑现。

**社会许可与宗教关切。** 博彩在斯里兰卡仍具敏感性。City of Dreams 开业证明政策可行，并不自动转化为第二个综合体的社会默许。项目将本地居民入场限制、问题赌博干预与社区基金作为与牌照申请同步的承诺，而非附属公关。

**先发者与市场教育。** City of Dreams 已占用“南亚首个综合度假村”心智。本项目的机会在错位（UHNW、峰会、医养、金库），代价是必须独立完成超高净值客群的目的地教育，获客成本在 Y1–Y2 高于稳态。爬坡修正情景（6.2.2）即为此设立。
'''

V11_443 = r'''
#### 弹性租金条款【V1.1 新增，并入 4.4.3】

V1.0 以 $1,000 万/厅/年作为招标厅定价锚。该锚仍是**稳态合约面值**。为提高开业前签约率、降低运营商断约概率，并与 7.2.2 节预售融资衔接，V1.1 增加弹性租金机制。弹性条款改变租金的时间分布与触发条件，**不改变 12 厅 × $1,000 万 = $1.20 亿的稳态年化口径**。

**1. 爬坡年折扣（仅限开业后前 24 个月）**

| 期间 | 计租规则 | 12 厅合计（若全部适用） | 说明 |
| :--- | :--- | :--- | :--- |
| 开业后第 1–12 个月 | 面值的 70% | 约 $8,400 万 | 与 6.9.4 节 Y1 部分厅年中开业、约 $4,200 万–$5,000 万入账的口径相容 |
| 开业后第 13–24 个月 | 面值的 85% | 约 $10,200 万 | 对应 Y2 上座率 55–62% 的过渡 |
| 第 25 个月起 | 面值 100% | $12,000 万 | 进入 6.2.2 基准情景 |

先锋租户若按 7.2.2 节预付两年租金，可在上述折扣上再享 5–10% 优惠，但预付额不得低于面值两年租金的 80%，差额以保证金形式托管。

**2. 保底 + 分成（可选，替代纯固定租）**

对客源尚未完全锁定的运营商，可选择：

- **保底租金**：不低于面值的 60%（即 $600 万/厅/年）；
- **GGR 分成**：厅内 GGR 超过约定门槛后，公司额外提取 8–12%；
- **合计上限**：保底 + 分成以面值 120% 为上限，避免景气年过度抽成伤害续约。

公司财务模型的基准情景仍按纯固定 $1,000 万/厅编制；保底 + 分成仅作为招商备选，不用于抬高呈报利润。

**3. 履约保障**

- 所有预售与租金预付进入 CPCEC 认可国际托管账户，项目未按期开业则退款并计违约金（沿用 7.2.2）。
- 单厅履约保证金不低于 $200 万，可与合作厅入场保证金规则对齐。
- 连续两期欠租：公司有权收回经营权并在剩余租期内重新招标；保证金优先弥补空置损失。
- 不可抗力（口岸长期关闭、牌照被暂停等）允许租金递延最长 6 个月，递延部分在后续 12 个月内摊还，不自动豁免。

**4. 与安全垫的关系**

1.6 节保守情景中的 $1.20 亿租金安全垫，在 V1.1 中应理解为**稳态契约面值**。Y1 若适用爬坡折扣，固定租金入账低于面值，这已反映在 6.9.4 与 6.2.2 爬坡修正情景中，不再把 Y1 按 $1.20 亿计入。
'''

V11_538 = r'''
### 5.3.8 客群达成前提【V1.1 新增】

第 5.3.1–5.3.7 节给出的客群占比与泥码区间，是**目标结构**，不是开业自动实现的结果。下列前提若缺失，对应客群份额应下调，并改用 6.2.2 爬坡修正情景，而非继续引用基准 70% 上座率。

| 客群 | 目标占比 | 必须具备的外部前提 | 项目侧可交付条件 | 未达成时的修正 |
| :--- | :--- | :--- | :--- | :--- |
| 海湾超高净值 | 30–35% | 迪拜/多哈/利雅得至科伦坡稳定航线；公务机与签证便利；清真供应链可认证 | 阿语礼宾、清真餐饮、家族出行与保管箱 | 占比下调至 15–20%，泥码中枢下移 |
| 金砖政商精英 | 20–25% | 峰会档期纳入外交与商务日历；主要成员国往来签证可预期 | 国宾楼层与 14 国主题包厢按期交付 | 会议引流延后，Y1–Y2 以酒店+医养为主 |
| 印度超富裕 | 20–25% | 孟买/德里航班密度；支付与 FEMA 合规通道清晰 | 印地语/英语礼宾，医疗旅游转介 | 以休闲酒店客替代，主大厅占比上升、VIP 泥码下降 |
| 东亚 VIP | 10–15% | 贵宾代理可在 GRA 框架下合规展业；跨境支付可审计 | 代理认证、返佣与 AML 系统 | Y1 东亚 VIP 按修正情景 1/3 计入 |
| 医养客群 | 8–12% | 再生医学/细胞治疗等项目取得可执业许可 | 右翼 2–6 层按期竣工、国际医师执业安排 | 康养收入按 6.6 节 50% 计入 |
| 会议外交 | 5–8% | 至少 1 场政府背书的国际会议落户 | 同传、安防、国宾动线验收 | 不影响赌场安全垫，但削弱品牌外溢 |

**管理规则（V1.1）：** 开业前 12 个月编制《客群前提清单》，每季度由项目公司向 CPCEC/GRA 更新一次红黄绿灯。若海湾 + 印度合计绿灯不足，则启动修正情景，并动用 7.2.1 风险准备金中的市场投放额度，而不是上调基准利润预测。
'''

V11_622 = r'''
#### 爬坡修正情景【V1.1 新增，并入 6.2.2】

V1.0 的 6.2.2 节给出保守 / 基准 / 乐观三档**稳态（约第 3 年）**参数。6.9.4 节已用文字说明 Y1 约 40%、Y2 约 75% 的爬坡。V1.1 将爬坡量化为一张可审计的修正表，供政府与融资方在牌照审查、贷款承诺和招标厅招商时使用。

**修正情景不替代基准情景。** 呈报利润、IRR 与第八章财政贡献的主表仍用基准稳态。修正情景用于：（1）Y1–Y2 现金流与偿债覆盖；（2）招标厅弹性租金测算；（3）当 5.3.8 前提未亮灯时的内部管理预算。

VIP 厅爬坡参数（26 厅合计逻辑；泥码与胜率沿用 6.2.2：净留存率 1.60%）

| 指标 | Y1 修正 | Y2 修正 | Y3 基准（V1.0） | Y3 修正（前提部分未达成） |
| :--- | :--- | :--- | :--- | :--- |
| 平均上座率 | 38% | 58% | 70% | 55% |
| 客均全程泥码 | $150,000 | $200,000 | $250,000 | $180,000 |
| 单厅年 GGR（约） | $480 万级 | $1,000 万级 | $1,509 万 | $900 万级 |
| 招标厅计租 | 面值 70% × 约 10 厅 × 约 5–8 个月 | 面值 85–100% × 12 厅全年 | 面值 $12,000 万 | 保底 60–85% 或 10 厅在租 |
| 公司确认赌场营收（约） | $1.10 亿（与 6.9.3 一致） | $2.90 亿 | $3.81 亿 | $2.4–2.8 亿 |

**与 7.4.2 保守情景的衔接：** 若 Y3 仍停留在上座率 50%、客均泥码 $150,000，则回到 V1.0 已披露的保守利润（净利润约 $4,200 万），前提是招标厅租金按弹性条款仍有保底流入。若同时出现大规模退租，则动用保证金 + 7.2.1 风险准备金覆盖 6–12 个月空置，不在未获新租约前继续宣称“固定 $1.20 亿”。

**主大厅：** 修正情景下有效运营率由 48% 下调至 Y1 的 28–32%、Y2 的 38–42%，Y3 回到 48%。主大厅不是安全垫来源。

**使用纪律：** 对外融资模型可同时附基准与修正两套；对政府财政贡献主表继续用基准，并在 8.2.6 给出政策与爬坡双维度敏感性，避免只报送上行数字。
'''

V11_721 = r'''
#### 风险准备金【V1.1 新增，并入 7.2.1】

V1.0 融资五渠合计 $9.50 亿不变。V1.1 在不扩大总投资的前提下，将已有缓冲**显式环圈**为风险准备金，防止弹性融资被提前用于扩装或提前分红。

| 来源 | 金额 | 环圈规则 |
| :--- | :--- | :--- |
| 建设应急储备（7.1.2 ⑩） | $3,100 万 | 仅用于造价 / 工期超支，动用须董事会 + 监理联签 |
| 弹性融资渠道中的专项额度（7.2.1 ⑤ 的一部分） | $5,000 万 | 从 $1.05 亿弹性渠道中环圈；未动用不得视为可分配现金 |
| **风险准备金合计** | **$8,100 万** | 约占总投资 8.5% |

允许动用的四类用途（须事先设定触发指标）：

1. **建设超支与关键路径突击**：对应 11.7.3 方案 A/C，单次不超过应急储备余额。
2. **开业延迟的营运资本**：酒店已开业而赌场牌照滞后超过 3 个月时，覆盖固定成本与债务利息，最长 12 个月。
3. **招标厅空置或折扣**：Y1–Y2 因弹性租金或退租导致固定租金低于面值的差额，优先以保证金弥补，不足部分由准备金补足至偿债覆盖率约定值。
4. **客群前提未亮灯时的市场投放**：5.3.8 项下海湾 / 印度航线或代理网络滞后，用于合规渠道建设，而非提高返佣突破 GRA 上限。

**明确不得动用：** 股东分红、非关键设计变更、非牌照必需的扩建、关联方借款。

**与 Category C 的关系：** 若政府与投资方共同选择将总投资增至 $10.00 亿以申报 Primary BSI Category C，新增 $5,000 万须来自主投资方追加或新的股权，**不得挤占**上述 $8,100 万准备金。Category C 路径见 9.5.1。

**披露原则：** 贷款协议与授权人协议应把准备金余额列为定期报告科目。V1.1 财务主表（7.4）仍按原 $9.50 亿与 20% 所得税编制，不因设立准备金而调高利润。
'''

V11_826 = r'''
#### 政策敏感性分析【V1.1 新增，并入 8.2.6】

上表十年累计约 $17.6 亿为**基准政策组合**下的预测：博彩 GGR 征收税 18%、企业所得税 20%、牌照费与入场税按 V1.0。以下敏感性用于回答“若税制或爬坡与呈报不一致，财政贡献如何变化”。数字为数量级测算，便于审阅，不构成新的承诺。

**A. 博彩 GGR 税率（以 Y3 全口径 GGR $4.42 亿为底）**

| GGR 税率 | Y3 GGR 税（万 USD） | 相对基准 |
| :--- | :--- | :--- |
| 15% | 约 6,630 | −16.7% |
| **18%（呈报基准）** | **7,964** | — |
| 22% | 约 9,730 | +22.2% |

税率上调增加国库收入、压缩运营商与招标厅承租意愿；V1.1 弹性租金与保底条款用于吸收部分冲击，但不能抵消税率跃升对 VIP 泥码的需求抑制。

**B. 企业所得税路径（以 Y3 EBT $21,654 万为底）**

| 路径 | 有效 CIT | Y3 企业所得税（万 USD） | 说明 |
| :--- | :--- | :--- | :--- |
| 本岛标准 | 30% | 约 6,496 | 未获港口城优惠时的下限情景 |
| **授权人协议 20%（呈报基准）** | **20%** | **4,331** | V1.0 / V1.1 主表 |
| Secondary BSI 4 年 7.5% 之后恢复 | 7.5%（前 4 年运营） | 前 4 年显著低于基准 | 2025 年 BSI 条例下的次级路径 |
| Primary BSI Category B 免税期 | 0%（免税年限内） | 0 | 见 9.5.1；免税期满后恢复当时税法 |
| Primary BSI Category C 免税期 | 0%（免税年限更长） | 0 | 需总投资达标 $10 亿 |

**对政府的含义：** 基准表按 20% CIT **多报了**免税期情景下的所得税、**少报了**未获优惠时的所得税。V1.1 要求在授权人协议谈判时锁定路径，并在十年财政表中以脚注标明“基准 = 20% CIT，不提前计入免税”。若最终获得 Category B/C 免税，所得税项下调、GGR 税仍按 GRA 规则征收——**GGR 税才是对国库更稳定的贡献**，不依赖 BSI 称号。

**C. 开业延迟 12 个月**

Y1 财政贡献约 $4,500 万整体后移一年；十年累计因折现与爬坡重叠，大约减少 0.4–0.8 成，取决于牌照是否在酒店已开业期间仍能收取部分非博彩税。延迟不自动取消十年 $17.6 亿量级，但改变“何时入库”。

**D. 爬坡修正 vs 基准（财政）**

若 Y3 停留在 6.2.2 修正情景（上座率约 55%、泥码 $180,000），GGR 税与所得税大致落在基准的 55–70% 区间，仍显著高于零，且招标厅保底租金继续贡献可预测的牌照费与部分税基。这与“项目失败则财政贡献归零”不是同一量级。
'''

V11_951 = r'''
#### Category C 及其他 BSI 税务路径【V1.1 新增，并入 9.5.1】

V1.0 以企业所得税 **20%** 作为授权人协议中的合同优惠税率，并据此编制第七、八章。V1.1 补充 2025 年《科伦坡港口城（对具有战略重要性业务给予豁免或激励的指引）条例》下 Primary / Secondary Business of Strategic Importance（BSI）分档，避免把“可能的免税期”写成已经到手的数字。

**Primary BSI 分档（条例第 1 号 / 2025，公开摘要）**

| 类别 | 最低投资（每地块或按细分地块比例） | 最低就业 | 最长项目实施期 | 实施期结束后 CIT 豁免年限 |
| :--- | :--- | :--- | :--- | :--- |
| A | $1.00 亿 | 300 人 | 5 年 | 10 年 |
| B | $5.00 亿 | 300 人 | 6 年 | 12 年 |
| C | **$10.00 亿** | 300 人 | 8 年 | **15 年** |
| D | $0.25 亿（Marina / 社会基础设施） | 100 人 | 4 年 | 8 年 |

本项目：总投资 **$9.50 亿**、直接就业 **4,100 人**、建设期约 3 年。就业与工期满足 A/B/C 的就业与实施期门槛；**投资额达到 Category B，距 Category C 还差 $5,000 万。**

**V1.1 税务策略（呈报口径优先稳健）**

1. **主表不变：** 继续按 20% CIT 建模。不把 12 或 15 年免税写入 IRR、NPV、八年分红与第八章主表。若日后获批免税，属于对股东的上行、对第八章所得税项的下调，另行通报政府。
2. **申报路径：** 以 Primary BSI **Category B** 为可立即满足的申请档；Category C 作为可选升级，须主投资方追加 $5,000 万真金白银（不得挪用 7.2.1 风险准备金）。
3. **Secondary BSI：** 若未能指定为 Primary，仍可能适用商业运营起 4 年 7.5% 优惠。该路径优于本岛 30%，仍优于或接近主表 20% 的前四年，具体以公报与授权人协议为准。
4. **与博彩税的关系：** Primary BSI 可获《博彩与投注征费法》（1988 年第 40 号）等列明法例的豁免，**不等于**自动免除 2025 年《博彩管理局法》下 GRA 对 GGR 的征收。本计划书第八章 18% GGR 税维持为对国库的核心贡献假设，除非 GRA 与财政部另有书面规则。
5. **实施期内进口环节：** 条例对 Primary BSI 在项目实施期内给予关税、港口与机场发展税、出口发展税等豁免。建设期造价已按远大 EPC 效率编制，进口豁免若落实，体现为应急储备动用概率下降，不用于上调利润预测。

**向政府的明确请求：** 请 CPCEC 在综合授权谈判中书面确认：（a）本项目作为授权人可适用的 BSI 类别；（b）CIT 豁免起算点（实施期结束 vs 商业运营）；（c）GRA GGR 税与 1988 年征费法豁免的适用边界。确认前，任何对外材料不得宣称“15 年免税已获批”。
'''

V11_1163 = r'''
#### 贵宾代理培育专项【V1.1 新增，并入 11.6.3】

Y2 的规模化（员工满编、代理人数、合规年审）以 Y1 代理网络真正跑通为前提。V1.0 将“贵宾代理网络快速激活”列为 Y1 举措；V1.1 升格为与土建并行的**专项计划**，预算从销售费用与 7.2.1 市场投放额度列支，不另增总投资。

**目标头寸**

| 时点 | 活跃持证代理 | 覆盖城市（优先） | 对应客群前提（5.3.8） |
| :--- | :--- | :--- | :--- |
| 开业前 12 个月 | 完成制度、合同、AML 培训大纲 | — | 红黄灯清单上线 |
| 开业后 90 天 | 40–60 人完成认证 | 澳门、香港、迪拜、孟买 | 东亚 + 海湾 + 印度启动 |
| Y1 年末 | 约 80–100 人 | 增加多哈、利雅得、新加坡 | 若航线未亮灯则暂缓海湾扩编 |
| Y2 年末 | ≥150 人（与 6.10 KPI 对齐） | 再增莫斯科/圣保罗等金砖节点 | 与峰会日历联动 |
| Y3 | 维持 ≥150，淘汰不合规者 | 以质量代数量 | 上座率目标 70% |

**四项不可省略的合规动作**

1. **GRA 框架内展业：** 代理不经手客户资金、不参与赌场运营决策（与 5.6.2 外部推广大使规则一致）；返佣进入可审计账户。
2. **FIT & PROPER：** 背景调查、制裁名单筛查、利益冲突申报；未通过者不得接待泥码客户。
3. **培训节奏：** 开业前 24 个月启动荷官/监察国际培训的同时，代理接受负责任博彩、AML、数据保护专项课，课时与考核记录保存不少于 7 年。
4. **与弹性租金联动：** 招标厅承租方自带代理网络的，须并入同一认证名单，避免“厅内一套、公司一套”的监管套利。

**失败触发：** 开业后 180 天活跃代理 <30 人或任一核心城市（迪拜 / 孟买 / 港澳）为零，则自动切换 6.2.2 修正情景，并召开董事会检讨 5.3.8 红灯项，而不是上修 Y1 利润。
'''

APPENDIX = r'''# 附录：V1.0→V1.1修订摘要

本附录说明修订版做了什么、**没有**改什么。V1.1 的目的是把政府审阅中可预见的质疑前置写入正文，而不是改写项目规模或上调回报。

## 未改动的核心口径（保持 V1.0）

| 项目 | V1.0 / V1.1 共同口径 |
| :--- | :--- |
| 总投资 | 约 $9.50 亿 |
| 地块 | 港口城 1-01-13，永久产权 14,000 m² |
| 开业 | 2029 年第二季度梯次开业 |
| Y3 基准营收 / EBITDA / 净利润 | 约 $6.40 亿 / $2.82 亿 / $1.73 亿 |
| 招标厅稳态面值 | 12 × $1,000 万/年 = $1.20 亿 |
| 财政贡献主表 | Y3 约 $1.53 亿，十年约 $17.6 亿 |
| 企业所得税主表 | 20%（不提前计入 BSI 免税） |
| GGR 税主表 | 18% |

## 修订清单

| 章节 | 修订点 | 性质 | 对数字的影响 |
| :--- | :--- | :--- | :--- |
| 封面 | 版本升至 V1.1（2026年9月）；呈报对象增加 CPCEC | 程序 | 无 |
| 1.6A | 宏观风险补充披露 | 新增 | 限定 1.6 节“不亏损”的前提 |
| 2.1.4 | 结构性挑战 | 新增 | 无；约束实施叙事 |
| 4.4.3 | 弹性租金条款 | 新增并入 | Y1–Y2 租金按 70%/85% 计，稳态面值不变 |
| 5.3.8 | 客群达成前提 | 新增 | 未亮灯则改用修正情景 |
| 6.2.2 | 爬坡修正情景 | 新增并入 | 管理预算 / 偿债；主表仍用基准 |
| 7.2.1 | 风险准备金 $8,100 万环圈 | 新增并入 | 总投资不变；弹性渠道中 $5,000 万不可分红 |
| 8.2.6 | 政策敏感性（GGR 税率、CIT 路径、延期、爬坡） | 新增并入 | 十年 $17.6 亿主表保留，附情景 |
| 9.5.1 | Category C / BSI 分档与 $9.50 亿 vs $10 亿 | 新增并入 | 主表仍 20% CIT；C 档需追加投资 |
| 11.6.3 | 贵宾代理培育专项 | 新增并入 | 不增总投资；失败则切修正情景 |
| 本附录 | 修订对照 | 新增 | — |

## 给审阅人的三句话

1. **回报没有因为修订而提高。** 所有新条款都在收紧前提、增加弹性与准备金。
2. **固定租金仍是结构优势，但 Y1 不再按 $1.20 亿现金入账来承诺。** 稳态面值与爬坡折扣已经分开写清。
3. **Category C 的 15 年 CIT 豁免是可选升级，不是当前模型的输入。** 现行 $9.50 亿对应可申请的最高确定档为 Category B；是否追加 $5,000 万由投资方与 CPCEC 另行决定。

编制单位：金砖娱乐（斯里兰卡）投资有限公司  
文件版本：V1.1 · 2026年9月
'''

INSERTIONS = [
    ("1.7 对斯里兰卡政府的价值：五维贡献", V11_16A),
    ("2.2 科伦坡港口城：斯里兰卡最重要的战略棋子", V11_214),
    ("4.5 右翼楼——金砖国际公寓酒店暨南亚康养旗舰中心", V11_443),
    ("5.4 竞争格局分析", V11_538),
    ("6.2.3 赌场主大厅GGR测算", V11_622),
    ("7.2.2 VIP贵宾厅经营权预售详解", V11_721),
    ("8.3 外汇净流入贡献", V11_826),
    ("9.5.2 税务合规承诺", V11_951),
    ("11.7 关键路径与工期风险管理", V11_1163),
]


def to_markdown_line(text: str) -> str:
    # True chapter titles look like "第十一章 实施时间表…". Cross-refs such as
    # "第四章4.3.1节" must stay as body text or they split the document.
    if re.match(r"^第[一二三四五六七八九十]+章(?=\s|$)", text):
        return f"# {text}"
    if re.match(r"^\d+\.\d+\.\d+\s+[\u4e00-\u9fffA-Za-z（(]", text) and len(text) <= 80:
        return f"### {text}"
    if re.match(r"^\d+\.\d+\s+[\u4e00-\u9fffA-Za-z（(]", text) and len(text) <= 80:
        return f"## {text}"
    return text


def _oxml_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def table_to_markdown(tbl_elem) -> str:
    rows: list[list[str]] = []
    for tr in tbl_elem.iterchildren():
        if tr.tag.split("}")[-1] != "tr":
            continue
        cells: list[str] = []
        for tc in tr.iterchildren():
            if tc.tag.split("}")[-1] != "tc":
                continue
            raw = " ".join(_oxml_text(tc).split())
            cells.append(raw.replace("|", "\\|"))
        if cells:
            rows.append(cells)
    if not rows or all(not c for row in rows for c in row):
        return ""
    width = max(len(r) for r in rows)
    for r in rows:
        if len(r) < width:
            r.extend([""] * (width - len(r)))
    header = rows[0]
    # If first row is blank placeholders, keep it as header anyway
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def v10_docx_to_markdown() -> str:
    """Rebuild V1.0 body from the original Word file, keeping real tables."""
    doc = Document(str(DOCX_SRC))
    chunks: list[str] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = _oxml_text(child)
            if text:
                chunks.append(to_markdown_line(text))
        elif tag == "tbl":
            md = table_to_markdown(child)
            if md:
                chunks.append(md)
    text = "\n\n".join(chunks)
    idx = text.find("# 第一章")
    if idx < 0:
        idx = text.find("第一章")
    if idx >= 0:
        text = text[idx:]
        if not text.startswith("# "):
            text = to_markdown_line(text.split("\n", 1)[0]) + text[text.find("\n"):]
    return text


def build_body_markdown() -> str:
    text = v10_docx_to_markdown()
    for anchor, block in INSERTIONS:
        patterns = [
            f"## {anchor}",
            f"### {anchor}",
            f"# {anchor}",
            anchor,
        ]
        placed = False
        for pat in patterns:
            if pat in text:
                text = text.replace(pat, block.rstrip() + "\n\n" + pat, 1)
                placed = True
                break
        if not placed:
            raise SystemExit(f"Anchor not found: {anchor}")
    return text


def split_chapters(body: str) -> dict[str, str]:
    parts = re.split(r"(?m)^# 第", body)
    out = {}
    # first chunk may be empty
    chapters = []
    for chunk in parts:
        chunk = chunk.strip()
        if not chunk:
            continue
        if not chunk.startswith("一") and "章" not in chunk[:8]:
            # reconstructed
            chunk = "第" + chunk if False else chunk
        text = chunk if chunk.startswith("第") else "第" + chunk
        text = "# " + text if not text.startswith("# ") else text
        # fix double hash
        if text.startswith("# 第"):
            pass
        elif text.startswith("# #"):
            text = text[2:]
        chapters.append(text)
    # The split removed '# 第' so restore
    restored = []
    for chunk in parts:
        c = chunk.strip()
        if not c:
            continue
        restored.append("# 第" + c if not c.startswith("#") else c)
    if len(restored) < 11:
        # fallback split on headings already in body
        restored = []
        bits = re.split(r"(?m)^(# 第[一二三四五六七八九十]+章\s+.*)$", body)
        # bits: preamble, h1, body1, h2, body2...
        i = 0
        if bits and not bits[0].startswith("# 第"):
            i = 1
        while i < len(bits) - 1:
            restored.append(bits[i].strip() + "\n\n" + bits[i + 1].strip())
            i += 2
    mapping = {}
    for fname, title in CHAPTER_FILES:
        found = None
        for ch in restored:
            if title in ch.splitlines()[0]:
                found = ch.strip() + "\n"
                break
        if not found:
            raise SystemExit(f"Chapter not found: {title}")
        mapping[fname] = found
    return mapping


def write_split_files(cover: str, chapters: dict[str, str], appendix: str) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    (OUT_DIR / "00_封面与目录.md").write_text(cover.strip() + "\n", encoding="utf-8")
    for fname, content in chapters.items():
        (OUT_DIR / fname).write_text(content.rstrip() + "\n", encoding="utf-8")
    (OUT_DIR / "12_附录_修订摘要.md").write_text(appendix.strip() + "\n", encoding="utf-8")
    readme = """# 金砖娱乐城商业计划书 V1.1（分章）

本目录为修订版 V1.1 的分章 Markdown。阅读顺序从 `00` 至 `12`。

完整合并稿：

- `../business-plan/金砖娱乐城综合开发项目商业计划书_V1.1.md`
- `../business-plan/金砖娱乐城综合开发项目商业计划书_V1.1.docx`

【V1.1 新增】段落已嵌入对应章节，不是独立备忘录。
"""
    (OUT_DIR / "README.md").write_text(readme, encoding="utf-8")


def write_combined_md(cover: str, chapters: dict[str, str], appendix: str) -> str:
    COMBINED_MD.parent.mkdir(parents=True, exist_ok=True)
    pieces = [cover.strip(), ""]
    for fname, _title in CHAPTER_FILES:
        pieces.append(chapters[fname].rstrip())
        pieces.append("")
    pieces.append(appendix.strip())
    text = "\n\n".join(pieces).strip() + "\n"
    COMBINED_MD.write_text(text, encoding="utf-8")
    return text


def iter_md_blocks(md: str):
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        sep = lines[i + 1] if i + 1 < len(lines) else ""
        if line.startswith("|") and sep.startswith("|") and re.search(r"-{3,}", sep):
            # table
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(lines[i])
                i += 1
            yield ("table", rows)
            continue
        if line.startswith("# "):
            yield ("h1", line[2:].strip())
        elif line.startswith("## "):
            yield ("h2", line[3:].strip())
        elif line.startswith("### "):
            yield ("h3", line[4:].strip())
        elif line.startswith("#### "):
            yield ("h4", line[5:].strip())
        elif line.startswith("---"):
            yield ("break", "")
        elif line.strip() == "":
            yield ("empty", "")
        else:
            yield ("p", line)
        i += 1


def add_run_font(run, size=11, bold=False, color=None, east_asia="Microsoft YaHei"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), east_asia)


def parse_table(rows: list[str]) -> list[list[str]]:
    out = []
    for r in rows:
        if re.match(r"^\|[ :\-|]+\|$", r.replace(" ", "") if False else r):
            if set(r.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set():
                continue
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells):
            continue
        out.append(cells)
    return out


def write_docx(md: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("机密 · 仅供政府审阅  |  金砖娱乐城商业计划书 V1.1")
    add_run_font(r, size=8, color=RGBColor(0x8B, 0x1E, 0x1E))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Bullion Entertainment City  ·  V1.1  ·  2026年9月  ·  未经授权不得披露")
    add_run_font(r, size=8, color=RGBColor(0x66, 0x66, 0x66))

    navy = RGBColor(0x0B, 0x1F, 0x3A)
    gold = RGBColor(0x8A, 0x6D, 0x3B)

    for kind, payload in iter_md_blocks(md):
        if kind == "empty":
            continue
        if kind == "break":
            continue
        if kind == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(payload)
            add_run_font(run, size=18, bold=True, color=navy)
        elif kind == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(payload)
            add_run_font(run, size=14, bold=True, color=navy)
        elif kind == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(payload)
            add_run_font(run, size=12, bold=True, color=gold)
        elif kind == "h4":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(payload)
            add_run_font(run, size=11, bold=True, color=gold)
        elif kind == "p":
            text = payload
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            # bold markdown ** **
            parts = re.split(r"(\*\*.+?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    add_run_font(run, size=11, bold=True)
                elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                    run = p.add_run(part[1:-1])
                    add_run_font(run, size=11)
                    run.italic = True
                else:
                    run = p.add_run(part)
                    add_run_font(run, size=11)
        elif kind == "table":
            grid = parse_table(payload)
            if not grid:
                continue
            cols = max(len(r) for r in grid)
            table = doc.add_table(rows=len(grid), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(grid):
                for ci in range(cols):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    val = row[ci] if ci < len(row) else ""
                    run = p.add_run(val.replace("**", ""))
                    add_run_font(run, size=9, bold=(ri == 0))
            doc.add_paragraph()

    COMBINED_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(COMBINED_DOCX))


def main():
    body = build_body_markdown()
    chapters = split_chapters(body)
    write_split_files(COVER, chapters, APPENDIX)
    combined = write_combined_md(COVER, chapters, APPENDIX)
    write_docx(combined)
    print("chapters:", len(chapters))
    print("out_dir:", OUT_DIR)
    print("md:", COMBINED_MD, "bytes", COMBINED_MD.stat().st_size)
    print("docx:", COMBINED_DOCX, "bytes", COMBINED_DOCX.stat().st_size)
    for p in sorted(OUT_DIR.iterdir()):
        print(f"  {p.name:40s} {p.stat().st_size:8d}")


if __name__ == "__main__":
    main()
